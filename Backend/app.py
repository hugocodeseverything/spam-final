"""
NOTE COPYRIGHT OF HUGO SACHIO WIJAYA
DO NOT DISTRIBUTE WITHOUT EXPLICIT PERMISSION FROM AUTHOR
"""
from flask import Flask, request, jsonify
from flask_cors import CORS
import pickle
import os
import json
from docx import Document
from PyPDF2 import PdfReader
from datetime import datetime
app = Flask(__name__)
CORS(app)
with open("models/stacked_spam_model.pkl", "rb") as f:
    stacked_model = pickle.load(f)
with open("models/nb_model.pkl", "rb") as f:
    nb_model = pickle.load(f)
with open("models/rf_model.pkl", "rb") as f:
    rf_model = pickle.load(f)
def load_users():
    if not os.path.exists("users.json"):
        return {}
    with open("users.json", "r") as f:
        return json.load(f)
def save_users(users):
    with open("users.json", "w") as f:
        json.dump(users, f)
@app.route("/register", methods=["POST"])
def register():
    users = load_users()
    data = request.json
    username = data.get("username")
    password = data.get("password")
    if not username or not password:
        return jsonify({"error": "Username dan password wajib diisi."}), 400
    if username in users:
        return jsonify({"error": "Username sudah terdaftar."}), 409
    users[username] = password
    save_users(users)
    return jsonify({"message": "Registrasi berhasil."})
@app.route("/login", methods=["POST"])
def login():
    users = load_users()
    data = request.json
    username = data.get("username")
    password = data.get("password")
    if users.get(username) == password:
        return jsonify({"message": "Login berhasil."})
    return jsonify({"error": "Username atau password salah."}), 401
@app.route("/predict", methods=["POST"])
def predict():
    data = request.json
    text = data.get("text", "")
    username = data.get("username", "anonymous")
    if not text:
        return jsonify({"error": "Teks kosong"}), 400
    label = {0: "Bukan Spam", 1: "Spam"}
    prediction = stacked_model.predict([text])[0]
    confidence = max(stacked_model.predict_proba([text])[0])
    result = label[prediction]
    try:
        with open("history.json", "r", encoding="utf-8") as f:
            history = json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        history = {}
    if username not in history:
        history[username] = []
    history[username].insert(0, {
        "timestamp": datetime.now().strftime("%d/%m/%Y, %H:%M:%S"),
        "text": text,
        "result": result,
        "confidence": round(confidence * 100, 2)
    })
    with open("history.json", "w", encoding="utf-8") as f:
        json.dump(history, f, indent=2, ensure_ascii=False)
    return jsonify({
        "naive_bayes": result,
        "nb_confidence": round(confidence * 100, 2),
        "random_forest": result,
        "rf_confidence": round(confidence * 100, 2),
        "input": text
    })
@app.route("/predict-file", methods=["POST"])
def predict_file():
    file = request.files.get("file")
    if not file:
        return jsonify({"error": "Tidak ada file yang dimasukan."}), 400
    filename = file.filename.lower()
    try:
        if filename.endswith(".docx"):
            doc = Document(file)
            text = "".join([para.text for para in doc.paragraphs if para.text.strip()])
        elif filename.endswith(".pdf"):
            reader = PdfReader(file)
            text = "".join([page.extract_text() for page in reader.pages if page.extract_text()])
        elif filename.endswith(".txt"):
            text = file.read().decode("utf-8")
        else:
            return jsonify({"error": "Format file tidak support. Cuman bisa .docx, .pdf, atau .txt"}), 400
    except Exception as e:
        return jsonify({"error": f"Gagal baca file: {str(e)}"}), 500
    prediction = stacked_model.predict([text])[0]
    confidence = max(stacked_model.predict_proba([text])[0])
    label = {0: "Bukan Spam", 1: "Spam"}
    result = label[prediction]
    try:
        with open("history.json", "r", encoding="utf-8") as f:
            history = json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        history = {}
    if "anonymous" not in history:
        history["anonymous"] = []
    history["anonymous"].insert(0, {
        "timestamp": datetime.now().strftime("%d/%m/%Y, %H:%M:%S"),
        "text": text[:200],
        "result": result,
        "confidence": round(confidence * 100, 2)
    })
    with open("history.json", "w", encoding="utf-8") as f:
        json.dump(history, f, indent=2, ensure_ascii=False)
    return jsonify({
        "naive_bayes": result,
        "nb_confidence": round(confidence * 100, 2),
        "random_forest": result,
        "rf_confidence": round(confidence * 100, 2),
        "input": text[:200] + ("..." if len(text) > 200 else "")
    })
@app.route("/history", methods=["GET"])
def get_history():
    username = request.args.get("username", "anonymous")
    try:
        with open("history.json", "r", encoding="utf-8") as f:
            history = json.load(f)
        return jsonify(history.get(username, []))
    except (FileNotFoundError, json.JSONDecodeError):
        return jsonify([])
if __name__ == "__main__":
    app.run(debug=True)
from werkzeug.utils import secure_filename
from docx import Document
from PyPDF2 import PdfReader
@app.route("/upload", methods=["POST"])
def upload_file():
    file = request.files.get("file")
    username = request.form.get("username", "anonymous")
    if not file:
        return jsonify({"error": "No file uploaded."}), 400
    filename = secure_filename(file.filename)
    ext = filename.split(".")[-1].lower()
    try:
        if ext == "pdf":
            reader = PdfReader(file)
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
        elif ext == "docx":
            doc = Document(file)
            text = "\n".join(p.text for p in doc.paragraphs)
        else:
            return jsonify({"error": "Unsupported file type."}), 400
    except Exception as e:
        return jsonify({"error": f"Failed to extract text: {str(e)}"}), 500
    if not text.strip():
        return jsonify({"error": "Text extraction failed or file is empty."}), 400

    label = {0: "Bukan Spam", 1: "Spam"}
    prediction = stacked_model.predict([text])[0]
    confidence = max(stacked_model.predict_proba([text])[0])
    result = label[prediction]
    try:
        with open("history.json", "r") as f:
            history = json.load(f)
    except:
        history = {}
    history.setdefault(username, []).append({
        "filename": filename,
        "text": text[:1000],
        "result": result,
        "confidence": confidence,
        "timestamp": datetime.now().isoformat()
    })
    with open("history.json", "w") as f:
        json.dump(history, f, indent=2)
    return jsonify({"result": result, "confidence": round(confidence * 100, 2)})
